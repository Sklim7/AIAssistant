import tkinter as tk
from tkinter import scrolledtext, ttk, filedialog
from docx import Document
from datetime import datetime
from openai import OpenAI
import threading
import os
import base64
from PIL import Image, ImageTk
import requests
from io import BytesIO
import sv_ttk  # Sun Valley ttk theme
import time

# AI Assistant Program

class MultiAIApp:
    def __init__(self, root):
        # app's UI, features
        self.root = root
        self.root.title("AI Assistant")
        self.root.geometry("1000x800")

        # Theme
        sv_ttk.set_theme("dark")  # dark theme

        # style settings
        self.style = ttk.Style()
        self.style.configure('Chat.TFrame', padding=10)
        self.style.configure('Control.TFrame', padding=5)
        self.style.configure('Title.TLabel', font=('Calibri', 15, 'bold'))
        self.style.configure('Status.TLabel', font=('Calibri', 12))

        # OpenAI client conversation history reset
        api_key = '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'
        self.client = OpenAI(api_key=api_key)
        self.conversation_history = []

        # UI Reset
        self.current_photo = None  # reference
        self.chat_area = None  # chat section
        self.message_entry = None  # message input section
        self.loading_label = None  # status label
        self.image_prompt = None  # image section prompt
        self.image_label = None  # created image section label
        self.image_status = None  # image status



        # calling UI section
        self.setup_ui()

    def setup_ui(self):
        # UI reset, arrange
        main_container = ttk.Frame(self.root, padding="10")
        main_container.pack(fill='both', expand=True)

        # notebook tab create
        self.notebook = ttk.Notebook(main_container)
        self.notebook.pack(fill='both', expand=True)

        # tab settings
        self.setup_chat_tab()  # chat tab
        self.setup_image_tab()  # image tab



    def setup_chat_tab(self):
        # chat
        self.chat_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.chat_frame, text=' 💬 Chat ')

        # Title Frame
        title_frame = ttk.Frame(self.chat_frame)
        title_frame.pack(fill='x', pady=(0, 10))
        ttk.Label(title_frame, text="Chat with AI", style='Title.TLabel').pack(side='left')

        # chat boundaries (outline and padding)
        chat_container = ttk.Frame(self.chat_frame, relief='solid', borderwidth=1)
        chat_container.pack(fill='both', expand=True, pady=(0, 10))

        self.chat_area = scrolledtext.ScrolledText(
            chat_container,             # widget in chat container
            wrap=tk.WORD,               # auto line enter
            height=20,                  # line spacing
            font=('Calibri', 15),       # Font: Calibri, size 10
            bg='#000000'                # background: black
        )
        self.chat_area.pack(fill='both', expand=True, padx=5, pady=5)

        # input container
        input_container = ttk.Frame(self.chat_frame, style='Control.TFrame')
        input_container.pack(fill='x', pady=(0, 5))

        # Message input frame
        input_frame = ttk.Frame(input_container)
        input_frame.pack(fill='x', pady=5)

        self.message_entry = ttk.Entry(
            input_frame,  # placing input at input frame
            font=('Calibri', 12)
        )
        self.message_entry.pack(
            side='left',
            fill='x',
            expand=True,
            padx=(0, 5)
        )
        self.message_entry.bind(
            '<Return>',  # Enter key
            self.send_message  # enterkey calls send message
        )

        ttk.Button(
            input_frame,
            text="Send",
            style='Accent.TButton',
            command=self.send_message

        ).pack(side='right')

        # Control Button frame
        control_frame = ttk.Frame(input_container)
        control_frame.pack(fill='x', pady=5)

        ttk.Button(
            control_frame,
            text="💾 Save Message",
            command=self.save_to_docx
        ).pack(side='left', padx=2)

        ttk.Button(
            control_frame,
            text="🔄 New Chat",
            command=self.new_chat
        ).pack(side='left', padx=2)


        # status display
        self.loading_label = ttk.Label(
            self.chat_frame,
            text="",
            style='Status.TLabel'
        )
        self.loading_label.pack(pady=5)

    def send_message(self, event=None):
        # sending message
        message = self.message_entry.get().strip()  # get text from input
        if not message:
            return

        # reset label user message
        self.message_entry.delete(0, tk.END)
        self.chat_area.insert(tk.END, f"User: {message}\n")
        self.chat_area.see(tk.END)

        # loading response
        self.loading_label.config(text="🤔 Loading Response...")
        threading.Thread(target=self.get_gpt_response, args=(message,)).start()

    def get_gpt_response(self, message):
        # Using OpenAI GPT
        try:
            self.conversation_history.append({"role": "user", "content": message})

            response = self.client.chat.completions.create(
                model="gpt-4o-mini",
                messages=self.conversation_history
            )

            assistant_message = response.choices[0].message.content  # Bring GPT resoponse
            self.conversation_history.append(
                {"role": "assistant", "content": assistant_message}
            )

            # UI update request
            self.root.after(0, self.update_chat_area, assistant_message)
        except Exception as e:
            self.root.after(0, self.update_chat_area, f"❌ Error: {str(e)}")
        finally:
            self.root.after(0, self.loading_label.config, {"text": ""})

    def update_chat_area(self, response_text):
        #Gpt response show
        self.chat_area.insert(tk.END, f"GPT: {response_text}\n\n")
        self.chat_area.see(tk.END)

    def new_chat(self):

        if self.chat_area.get("1.0", tk.END).strip():
            self.auto_save_chat()  # auto save last chat

        self.conversation_history = []  # chat history reset
        self.chat_area.delete("1.0", tk.END)
        self.chat_area.insert(tk.END, "✨ New chat has started!\n\n")
        self.chat_area.see(tk.END)

    def auto_save_chat(self):
        save_dir = "chat_history"
        if not os.path.exists(save_dir):
            os.makedirs(save_dir)

        filename = f"Chat_history_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        file_path = os.path.join(save_dir, filename)

        self.save_chat_to_file(file_path)
        self.chat_area.insert(tk.END, f"💾 Chat was saved: {file_path}\n\n")
        self.chat_area.see(tk.END)

    def save_chat_to_file(self, file_path):
        # saving chat as word file.
        doc = Document()
        doc.add_heading('AI Assistant chat history', 0)
        doc.add_paragraph(
            f"Timesaved: {datetime.now().strftime('%Y %m %d %H:%M:%S')}"
        )
        doc.add_paragraph('=' * 50)

        chat_content = self.chat_area.get("1.0", tk.END)
        paragraphs = chat_content.split('\n')

        for para in paragraphs:
            if para.strip():
                if para.startswith('Me:'):
                    p = doc.add_paragraph()
                    p.add_run('Me: ').bold = True
                    p.add_run(para[3:])
                elif para.startswith('GPT:'):
                    p = doc.add_paragraph()
                    p.add_run('GPT: ').bold = True
                    p.add_run(para[4:])
                else:
                    doc.add_paragraph(para)

        doc.save(file_path)

    def save_to_docx(self):

        try:
            file_path = filedialog.asksaveasfilename(
                defaultextension=".docx",
                filetypes=[("Word documents", "*.docx")],
                initialfile=f"Chathistory_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            )

            if not file_path:
                return

            self.save_chat_to_file(file_path)
            self.chat_area.insert(tk.END, f"💾 chat was saved: {file_path}\n\n")
            self.chat_area.see(tk.END)

        except Exception as e:
            self.chat_area.insert(tk.END,
                                  f"❌ Error: Error saving message. {str(e)}\n\n")
            self.chat_area.see(tk.END)


    def setup_image_tab(self):
        # Image tab Ui
        self.image_frame = ttk.Frame(self.notebook, style='Chat.TFrame')
        self.notebook.add(self.image_frame, text=' 🎨 Create Image ')

        # Title
        ttk.Label(
            self.image_frame,
            text="Create AI Image",
            style='Title.TLabel'
        ).pack(pady=(0, 10))


        control_panel = ttk.LabelFrame(self.image_frame, text="Image Settings", padding=10)
        control_panel.pack(fill='x', padx=5, pady=5)

        # Image_promt
        prompt_frame = ttk.Frame(control_panel)
        prompt_frame.pack(fill='x', pady=5)

        ttk.Label(prompt_frame, text="Image Description:").pack(side='left', padx=(0, 5))
        self.image_prompt = ttk.Entry(prompt_frame, width=50)
        self.image_prompt.pack(side='left', fill='x', expand=True)

        # Image settings
        settings_frame = ttk.Frame(control_panel)
        settings_frame.pack(fill='x', pady=10)

        # Image size
        size_frame = ttk.LabelFrame(settings_frame, text="Size", padding=5)
        size_frame.pack(side='left', padx=5)

        self.size_var = tk.StringVar(value="1024x1024")
        ttk.Radiobutton(
            size_frame,
            text="square",
            variable=self.size_var,
            value="1024x1024"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            size_frame,
            text="vertical",
            variable=self.size_var,
            value="1024x1792"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            size_frame,
            text="horizontal",
            variable=self.size_var,
            value="1792x1024"
        ).pack(side='left', padx=5)

        # quality selection
        quality_frame = ttk.LabelFrame(settings_frame, text="quality", padding=5)
        quality_frame.pack(side='left', padx=5)

        self.quality_var = tk.StringVar(value="standard")
        ttk.Radiobutton(
            quality_frame,
            text="standard",
            variable=self.quality_var,
            value="standard"
        ).pack(side='left', padx=5)

        ttk.Radiobutton(
            quality_frame,
            text="high quality",
            variable=self.quality_var,
            value="hd"
        ).pack(side='left', padx=5)

        # Create Button
        ttk.Button(
            control_panel,
            text="🎨 Create Image!",
            style='Accent.TButton',
            command=self.generate_image
        ).pack(pady=10)

        # Image display environment
        image_display = ttk.LabelFrame(self.image_frame, text="Created Image", padding=10)
        image_display.pack(fill='both', expand=True, padx=5, pady=5)

        self.image_label = ttk.Label(image_display)
        self.image_label.pack(pady=10)

        self.image_status = ttk.Label(
            image_display,
            text="",
            style='Status.TLabel'
        )
        self.image_status.pack(pady=5)



    def generate_image(self):
        # create image
        prompt = self.image_prompt.get().strip()
        if not prompt:
            self.image_status.config(text="⚠️ Please enter image description.")
            return

        self.image_status.config(text="🎨 Creating Image...")
        threading.Thread(target=self.generate_image_thread, args=(prompt,)).start()

    def generate_image_thread(self, prompt):
        try:
            # DALL-E API to create image
            response = self.client.images.generate(
                model="dall-e-3",  #  DALL-E model
                prompt=prompt,  # user input prompt
                size=self.size_var.get(),  # image size (selected from UI)
                quality=self.quality_var.get(),  # image quality (selected from UI)
                n=1,  # image quantity
                style="natural"  # image style
            )

            # get image URL
            image_url = response.data[0].url
            image_response = requests.get(image_url)
            image = Image.open(BytesIO(image_response.content))  # reading image from URl

            # image size  (max 800x800 pixel)
            max_size = (800, 800)
            image.thumbnail(max_size, Image.Resampling.LANCZOS)

            photo = ImageTk.PhotoImage(image)  # image for Tkinter

            # UI Image status
            self.root.after(0, self.update_image, photo, "✨ Image has been created!")

            # save image in local folder
            save_dir = "generated_images"  # directory
            if not os.path.exists(save_dir):
                os.makedirs(save_dir)  # create if no directory

            # add timestamp and save
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            save_path = os.path.join(save_dir, f"generated_image_{timestamp}.png")
            image.save(save_path)

        except Exception as e:
            # error case
            error_message = str(e)
            self.root.after(0, self.image_status.config,
                            {"text": f"❌ Error: {error_message}"})
            print(f"Error in image generation: {error_message}")

    def update_image(self, photo, status_text):
        # UI image update. status text update
        self.current_photo = photo  # image doesn't get deleted from TKinter
        self.image_label.config(image=photo)  # update image label
        self.image_status.config(text=status_text)  # updated status





if __name__ == "__main__":
    # initiate program
    root = tk.Tk()  # Tkinter window
    app = MultiAIApp(root)  # MultiAIApp class
    root.mainloop()  # Tkinter event loop start